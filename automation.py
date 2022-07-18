import openpyxl
from docx import Document

wb = openpyxl.load_workbook("employee_details.xlsx")
doc = Document("Letter_of_Training.docx")
sheet = wb["Sheet1"]
header = []
data = []
columns = sheet.max_column
# rows = sheet.max_row  #row count is wrong

# print("num of cols: ", columns)
# print("num of rows: ", rows)

for emp in range(2, 11):
    for cols in range(1, columns + 1):
        header.append(sheet.cell(1, cols).value)

    for cols in range(1, columns + 1):
        data.append(sheet.cell(emp, cols).value)
    #     print(cols)
    #     print(header)
    #     print(data)

    emp_dict = dict(zip(header, map(str, data)))
    #     print(emp_dict)
    #     print(" ")
    texts = []
    #     print(emp_dict["employee_name"])

    for p in doc.paragraphs:
        #     if p.text.isspace() == False:
        texts.append(p.text)
    # print(texts)
    newDoc = ""
    for sentense in texts:
        # print(sentense)
        for i in (("[employee_name]", emp_dict["employee_name"]), ("[department]", emp_dict["department"]),
                  ("[stipend]", emp_dict["stipend"]), ("[designation]", emp_dict["designation"]),
                  ("[joining_date]", emp_dict["date_of_joining "]), ("[email]", emp_dict["email"]),
                  ("[trainer_name]", emp_dict["trainer_name"])):
            sentense = sentense.replace(*i)

        newDoc += (sentense + "\n")
    print(newDoc)
    mydoc = Document()
    mydoc.add_paragraph(newDoc)

    if mydoc.save(f'{emp_dict["employee_name"]}_letter_of_trainning.docx'):
        print("new file generated")
    else:
        print("file generation failed")

    print("==================================================")
